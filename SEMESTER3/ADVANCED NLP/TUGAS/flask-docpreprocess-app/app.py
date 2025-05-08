from flask import Flask, render_template, request, redirect, url_for, flash, jsonify
import os
from werkzeug.utils import secure_filename
from datetime import datetime
import re
from collections import defaultdict
from docx import Document
import pandas as pd
from PyPDF2 import PdfReader
import odf.opendocument
from odf import text

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SECRET_KEY'] = 'your-secret-key'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Inverted index storage
inverted_index = defaultdict(dict)

def index_file(filename):
    """Index a file and add to inverted index"""
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    try:
        text = ""
        # Handle PDF
        if filename.lower().endswith('.pdf'):
            with open(filepath, 'rb') as f:
                pdf = PdfReader(f)
                text = " ".join([page.extract_text() for page in pdf.pages])
        
        # Handle Word DOCX
        elif filename.lower().endswith('.docx'):
            doc = Document(filepath)
            text = " ".join([para.text for para in doc.paragraphs])
        
        # Handle Legacy Word DOC (membutuhkan antiword)
        elif filename.lower().endswith('.doc'):
            try:
                text = os.popen(f'antiword "{filepath}"').read()
            except:
                text = ""
        
        # Handle Excel
        elif filename.lower().endswith(('.xlsx', '.xls')):
            df = pd.read_excel(filepath, sheet_name=None)
            text = " ".join(
                [str(cell) for sheet in df.values() 
                for row in sheet.values 
                for cell in row]
            )
        
        # Handle OpenDocument (ODT)
        elif filename.lower().endswith('.odt'):
            doc = odf.opendocument.load(filepath)
            text = " ".join(
                [text.Text(e).toString() 
                for e in doc.getElementsByType(text.P)]
            )
        
        # Handle plain text
        elif filename.lower().endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        # Bersihkan teks dari karakter khusus
        text = re.sub(r'\s+', ' ', text).strip()
        
        if not text:
            app.logger.warning(f"No text extracted from {filename}")
            return
        
        # Proses teks yang sudah diekstrak
        words = re.findall(r'\w+', text.lower())
        
        # Update inverted index
        for word in set(words):
            term_frequency = words.count(word)
            inverted_index[word][filename] = term_frequency

    except Exception as e:
        app.logger.error(f"Error processing {filename}: {str(e)}")
        flash(f"Error processing {filename}", "error")

def search_in_index(query, selected_files=None):
    """Search the inverted index for the query"""
    results = []
    query_terms = re.findall(r'\w+', query.lower())
    
    for term in query_terms:
        if term in inverted_index:
            for filename, term_frequency in inverted_index[term].items():
                # If specific files are selected, only search in those
                if selected_files and filename not in selected_files:
                    continue
                
                # Get snippets of text around the search term
                snippets = get_snippets(filename, term)
                
                # Add to results
                results.append({
                    'filename': filename,
                    'term': term,
                    'term_frequency': term_frequency,
                    'snippets': snippets
                })
    
    return results

def get_snippets(filename, term):
    """Get snippets of text around the search term (support multi-format)"""
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    text = ""
    
    try:
        # Replikasi logika ekstraksi teks dari index_file()
        if filename.lower().endswith('.pdf'):
            with open(filepath, 'rb') as f:
                pdf = PdfReader(f)
                text = " ".join([page.extract_text() for page in pdf.pages])
        
        elif filename.lower().endswith('.docx'):
            doc = Document(filepath)
            text = " ".join([para.text for para in doc.paragraphs])
        
        elif filename.lower().endswith('.xlsx'):
            df = pd.read_excel(filepath, sheet_name=None)
            text = " ".join([str(cell) for sheet in df.values() for row in sheet.values for cell in row])
        
        else:  # Untuk format teks biasa
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        # Cari snippet dengan konteks
        pattern = re.compile(r'(\b\w+\W+){0,5}\b' + re.escape(term) + r'\b(\W+\w+\b){0,5}', re.IGNORECASE)
        matches = pattern.finditer(text)
        
        snippets = []
        for match in matches:
            snippet = match.group(0)
            highlighted = snippet.replace(term, f'<span class="bg-yellow-200">{term}</span>')
            snippets.append(highlighted)
            if len(snippets) >= 3:
                break
        
        return snippets
    
    except Exception as e:
        app.logger.error(f"Error getting snippets from {filename}: {str(e)}")
        return []
    
@app.route('/')
def index():
    return render_template("index.html")

@app.route('/preprocess')
def preprocess():
    return render_template("prepocesing.html")

@app.route('/model')
def model():
    return render_template("model.html")

@app.route('/search_menu', methods=['GET', 'POST'])
def search():
    # Handle search request
    if request.method == 'POST':
        query = request.form.get('query', '').strip()
        selected_files = request.form.get('selected_files', '')
        selected_files = selected_files.split(',') if selected_files else None

        app.logger.debug(f"selected_files: {selected_files}")

        if not query :
            flash('Please enter a search term', 'error')
            return redirect(url_for('search'))

        # Validasi untuk selected_files
        if not selected_files or all(file == '' for file in selected_files):
            flash('At least one file must be selected', 'error')
            return redirect(url_for('search'))
                            
        search_results = search_in_index(query, selected_files)
    else:
        search_results = None
        query = None
    
    # Get uploaded files list
    uploaded_files = []
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.isfile(file_path):
            mtime = os.path.getmtime(file_path)
            uploaded_files.append({
                'name': filename,
                'date': datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M')
            })
    # Sort by upload time descending
    uploaded_files.sort(key=lambda x: x['date'], reverse=True)
    
    return render_template(
        'search.html', 
        uploaded_files=uploaded_files,
        search_results=search_results,
        search_query=query
    )

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'files' in request.files:
        files = request.files.getlist('files')
        uploaded_count = 0
        
        for file in files:
            if file.filename != '' and file:
                filename = secure_filename(file.filename)
                # Filter ekstensi file
                if not filename.lower().endswith(('.pdf', '.doc', '.docx', '.txt', '.xls', '.xlsx')):
                    continue
                
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                index_file(filename)
                uploaded_count += 1
        
        if uploaded_count > 0:
            flash(f'{uploaded_count} file(s) uploaded and indexed', 'success')
        else:
            flash('No valid files uploaded', 'error')
    
    return jsonify({'status': 'success'}), 200

@app.route('/delete/<filename>', methods=['POST'])
def delete_file(filename):
    # safe_filename = secure_filename(filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    app.logger.debug(f"Mencoba menghapus file: {file_path}")
    app.logger.debug(f"Filename asli: {filename}")
    app.logger.debug(f"Filename aman: {filename}")
    
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            # Remove from inverted index
            for term in list(inverted_index.keys()):
                if filename in inverted_index[term]:
                    del inverted_index[term][filename]
                    if not inverted_index[term]:  # Remove term if no more documents
                        del inverted_index[term]
            app.logger.debug(f"File {file_path} berhasil dihapus")
            flash('File berhasil dihapus', 'success')
        else:
            app.logger.warning(f"File {file_path} tidak ditemukan")
            flash('File tidak ditemukan', 'error')
    except Exception as e:
        app.logger.error(f"Gagal menghapus file: {str(e)}", exc_info=True)
        flash(f'Gagal menghapus file: {str(e)}', 'error')
    
    return redirect(url_for('search'))

if __name__ == '__main__':
    app.run(debug=True)